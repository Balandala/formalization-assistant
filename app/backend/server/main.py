import os
import uuid
import aiofiles
import requests
import httpx
from docxcompose.composer import Composer
from contextlib import asynccontextmanager
from fastapi import FastAPI, UploadFile, File, Form, Depends, HTTPException, APIRouter
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update
from fastapi.background import BackgroundTasks
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from docx import Document as DocxDocument
import tempfile

from backend.server.database import engine, Base, get_db
from backend.server.models import Document, TaskStatus
from backend.server.schemas import DocumentResponse
from shared.models import TitleData

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
FRONTEND_DIR = os.path.join(BASE_DIR, "frontend")
FRONTEND_STATIC_DIR = os.path.join(FRONTEND_DIR, "static")


@asynccontextmanager
async def lifespan(app: FastAPI):
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield

app = FastAPI(title="DocHelper", lifespan=lifespan)
router = APIRouter()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:8000",
        "http://127.0.0.1:8000",
        "http://localhost:7272",
        "http://127.0.0.1:7272",
        "http://localhost:7777",
        "http://127.0.0.1:7777",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.mount("/styles", StaticFiles(directory=os.path.join(FRONTEND_DIR, "styles")), name="styles")
app.mount("/scripts", StaticFiles(directory=os.path.join(FRONTEND_DIR, "scripts")), name="scripts")
app.mount("/static", StaticFiles(directory=os.path.join(FRONTEND_DIR, "static")), name="static")

@app.post("/upload", response_model=DocumentResponse)
async def upload_document(file: UploadFile = File(...),
                          db: AsyncSession = Depends(get_db)):

    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Not a docx")

    doc_id = uuid.uuid4()
    filename = f"{doc_id}_{file.filename}"
    file_path = os.path.join(UPLOAD_FOLDER, filename)

    async with aiofiles.open(file_path, 'wb') as out_file:
        while content := await file.read(1024*1024):
            await out_file.write(content)

    new_doc = Document(
        id=doc_id,
        filename=file.filename,
        path=file_path,
        status=TaskStatus.PENDING.value
    )
    db.add(new_doc)
    await db.commit()
    await db.refresh(new_doc)

    await process_doc(file_path, new_doc.id, db)

    return new_doc

@app.get("/status/{doc_id}", response_model=DocumentResponse)
async def get_status(doc_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    return doc

@app.get("/download/{doc_id}")
async def get_document(doc_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if doc.status != TaskStatus.COMPLETED.value:
        return HTTPException(status_code=400, detail="Document not ready")


    if not os.path.exists(doc.path):
        raise HTTPException(status_code=500, detail="Processed file not found on server")



    return FileResponse(
        path=doc.path,
        filename=doc.filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )

@app.get("/main")
async def get_main_page():
    filename = os.path.join(FRONTEND_STATIC_DIR, "index.html")
    if not os.path.exists(filename):
        raise HTTPException(status_code=404, detail="index.html not found")
    return FileResponse(filename)


async def process_doc(filepath: str, doc_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    URL = "https://127.0.0.1:7272/process"
    absPath = os.path.abspath(filepath)
    response = requests.post(URL, json={"filepath":absPath}, verify= False)
    if response.status_code == 200:
        stat = TaskStatus.COMPLETED.value
    else:
        stat = TaskStatus.FAILED.value
    stmt = update(Document).where(Document.id == doc_id).values(status=stat)
    await db.execute(stmt)
    await db.commit()


@router.post("/generate-title", response_model=DocumentResponse)
async def generate_title_endpoint(
    data: TitleData,
    db: AsyncSession = Depends(get_db)
):
    doc_id = uuid.uuid4()
    file_path = os.path.join(UPLOAD_FOLDER, f"{doc_id}.docx")

    async with httpx.AsyncClient(timeout=30.0) as client:
        try:
            response = await client.post(
                "http://127.0.0.1:7777/generate-title",
                json={
                    "doc_id": str(doc_id),
                    "data": data.model_dump()
                }
            )
            print(response)
            if response.status_code == 200:
                os.makedirs(UPLOAD_FOLDER, exist_ok=True)
                with open(file_path, "wb") as f:
                    f.write(response.content)

                new_doc = Document(
                    id=doc_id,
                    filename="title_page.docx",
                    path=file_path,
                    status=TaskStatus.COMPLETED.value
                )
                db.add(new_doc)
                await db.commit()
                await db.refresh(new_doc)
                return new_doc
            else:
                raise HTTPException(status_code=500, detail="Failed to generate title")
        except httpx.RequestError as exc:
            raise HTTPException(status_code=503, detail=f"Title service unreachable: {exc}")


@router.post("/upload-with-title", response_model=DocumentResponse)
async def upload_with_title(
    file: UploadFile = File(...),
    institute: str = Form(...),
    work_type: str = Form(...),
    subject: str = Form(...),
    theme: str = Form(...),
    author: str = Form(...),
    group: str = Form(...),
    chief: str = Form(...),
    post: str = Form(...),
    db: AsyncSession = Depends(get_db)
):
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files allowed")

    doc_id = uuid.uuid4()
    temp_input_path = os.path.join(UPLOAD_FOLDER, f"{doc_id}_input.docx")
    processed_path = os.path.join(UPLOAD_FOLDER, f"{doc_id}_processed.docx")
    title_path = os.path.join(UPLOAD_FOLDER, f"{doc_id}_title.docx")
    final_path = os.path.join(UPLOAD_FOLDER, f"{doc_id}_final.docx")


    try:
        async with aiofiles.open(temp_input_path, 'wb') as f:
            while chunk := await file.read(1024 * 1024):
                await f.write(chunk)


        URL = "https://127.0.0.1:7272/process"
        abs_input_path = os.path.abspath(temp_input_path)


        proc_response = requests.post(URL, json={"filepath": abs_input_path}, verify=False)

        if proc_response.status_code != 200:
            raise HTTPException(status_code=500, detail="Ошибка обработки файла микросервисом")


        processed_actual_path = temp_input_path


        title_data = TitleData(
            institute=institute, work_type=work_type, subject=subject,
            theme=theme, author=author, group=group, chief=chief, post=post
        )

        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(
                "http://127.0.0.1:7777/generate-title",
                json={"doc_id": str(doc_id), "data": title_data.model_dump()}
            )

        if resp.status_code != 200:
            raise HTTPException(status_code=500, detail="Ошибка генерации титульника")

        with open(title_path, 'wb') as f:
            f.write(resp.content)


        if not os.path.exists(processed_actual_path):
            raise HTTPException(status_code=500, detail="Обработанный файл исчез")

        merge_title_with_document(title_path, processed_actual_path, final_path)


        new_doc = Document(
            id=doc_id,
            filename=file.filename,
            path=final_path,  # Ссылка на итоговый файл
            status=TaskStatus.COMPLETED.value
        )

        db.add(new_doc)
        await db.commit()
        await db.refresh(new_doc)


        remove_file(temp_input_path)
        remove_file(title_path)


        return new_doc

    except Exception as e:
        # Очистка при ошибке
        for path in [temp_input_path, processed_path, title_path, final_path]:
            if os.path.exists(path):
                os.remove(path)
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

app.include_router(router)

def remove_file(path: str):
    if os.path.exists(path):
        os.remove(path)

def merge_title_with_document(title_path: str, content_path: str, output_path: str):

    master = DocxDocument(title_path)
    composer = Composer(master)

    doc_to_append = DocxDocument(content_path)

    composer.append(doc_to_append)

    composer.save(output_path)