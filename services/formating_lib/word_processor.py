from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


from formating_config import Config

import styles as styles_lib

import os
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class WordProcessor:
    def __init__(self, config: Config):
        self.config = config

    def process_file(self, filepath: str):
        if not filepath.endswith(".docx"):
            raise ValueError("Invalid file type. Only .docx is supported.")
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        try:
            logger.info(f"Opening document: {filepath}")
            doc = Document(filepath)
            logger.info("Document initialized successfully")


            self.process(doc)

            doc.save(filepath)
            logger.info("Document saved successfully")
        except Exception as e:
            logger.exception(f"Failed to process document {filepath} with Exception {e}")
            raise

    def process(self, doc):
        self._doc_init(doc)
        self._override_styles(doc)
        paragraphs = list(doc.paragraphs)

        for i, p in enumerate(paragraphs):
            if self.config.override_formatting:
                self._override_run_properties(p)

            if self.config.headings and self._is_title(p, paragraphs, i):
                p.style = styles_lib.StyleIds.Heading1
            elif self.config.captions and self._is_after_media(p, paragraphs, i):
                p.style = styles_lib.StyleIds.Media
            elif self._is_contains_media(p):
                p.style = styles_lib.StyleIds.Heading1
            elif self.config.normal_text:
                p.style = styles_lib.StyleIds.Normal

        if self.config.pages_numerations:
            self._add_footer(doc)
        if self.config.page_fields:
            self._add_page_margins(doc)

    def _doc_init(self, doc):
        if doc._element.body is None:
            body = OxmlElement("w:body")
            doc._element.append(body)

    def _override_styles(self, doc):
        # Это не функция это пиздец трогать только в антирадиционном костюме
        try:
            styles_part = doc.part.package.part_related_by(RT.STYLES)
        except KeyError:
            styles_part = None

        if styles_part is None:
            pkg = doc.part.package
            uri = "/word/styles.xml"
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"


            found_parts = [p for p in pkg.iter_parts() if p.partname == uri]
            if found_parts:
                styles_part = found_parts[0]
                styles_element = styles_part.element
            else:
                styles_xml = f'<w:styles {nsdecls("w")}></w:styles>'
                styles_element = parse_xml(styles_xml)
                styles_part = pkg._add_part(uri, content_type, styles_element)

            doc.part.relate_to(styles_part, RT.STYLES)
            logger.info("Created new styles.xml and relationship")
        else:
            styles_element = styles_part.element

        styles_element.clear()
        styles_element.append(styles_lib.make_text_style())
        styles_element.append(styles_lib.make_heading_style())
        styles_element.append(styles_lib.make_caption_style())

        if hasattr(doc, '_styles'):
            delattr(doc, '_styles')
        _ = doc.styles  # Принудительно перезагружается кэш стилей

    def _is_title(self, p, all_paragraphs, index):
        if not p.text.strip():
            return False

        score = 0

        if self._is_contains_media(p) or self._is_in_table(p):
            score -= 100

        if self._is_after_page_break(p, all_paragraphs, index):
            score += 2
        if self._is_text_short(p):
            score += 2
        if self._is_heading_by_structural_pattern(p):
            score += 2
        if self._is_heading_by_position(p, all_paragraphs, index):
            score += 2

        return score > 2

    def _is_in_table(self, p):
        parent = p._element.getparent()
        while parent is not None:
            if parent.tag.endswith("tbl"):
                return True
            parent = parent.getparent()
        return False

    def _is_after_page_break(self, p, all_paragraphs, index):
        if index == 0:
            return False
        prev_p = all_paragraphs[index - 1]
        for br in prev_p._element.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
        return False

    def _is_text_short(self, p):
        text = p.text.strip()
        return 0 < len(text) < 200

    def _is_heading_by_structural_pattern(self, p):
        text = p.text.strip()
        if not text:
            return False
        starts_with_number = text[0].isdigit()
        is_in_numbered_list = bool(p.style.name.startswith("List"))
        return starts_with_number and not is_in_numbered_list

    def _is_heading_by_position(self, p, all_paragraphs, index):
        if index == 0:
            return True
        if index > 0:
            prev_text = all_paragraphs[index - 1].text.strip()
            curr_text = p.text.strip()
            ends_with_punct = lambda t: t.endswith((".", "!", "?"))
            return ends_with_punct(prev_text) and not ends_with_punct(curr_text)
        return False

    def _is_after_media(self, p, all_paragraphs, index):
        if index == 0:
            return False
        prev_p = all_paragraphs[index - 1]
        if self._is_contains_media(prev_p):
            return len(p.text.strip()) < 200
        return False

    def _is_contains_media(self, p):
        for drawing in p._element.iter(qn("w:drawing")):
            if drawing.find(".//" + qn("pic:pic")) is not None:
                return True
        return False

    def _override_run_properties(self, p):
        for run in p.runs:
            run.bold = None
            run.italic = None
            run.underline = None
            run.font.color.rgb = None

    def _add_footer(self, doc):
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.clear()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

    def _add_page_margins(self, doc):
        section = doc.sections[0]
        sect_pr = section._sectPr
        pg_mar = sect_pr.find(qn("w:pgMar"))
        if pg_mar is None:
            pg_mar = OxmlElement("w:pgMar")
            sect_pr.append(pg_mar)
        pg_mar.set(qn("w:left"), "1700")
        pg_mar.set(qn("w:right"), "850")
        pg_mar.set(qn("w:top"), "1133")
        pg_mar.set(qn("w:bottom"), "1133")